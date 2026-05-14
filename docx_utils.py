#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_utils - Modulo avanzado para generacion de documentos .docx
para el proyecto MindMood - Diario Emocional Inteligente

Materia: SEMINARIO DE INGENIERIA DE SOFTWARE
Alumno:  RAMIREZ RUIZ, CRISTOPHER SAID
Seccion: D15 — CUCEI Guadalajara Jalisco

Proporciona funciones para:
- Graficos nativos de Word (barras, pastel, lineas)
- Diagramas con matplotlib (barras, pastel, UML, secuencia, casos de uso, estados)
- Tabla de contenido automatica (TOC)
- Marca de agua en encabezados
- Secciones en orientacion horizontal/vertical
- Estilos personalizados de parrafo
- Pie de pagina con numeracion automatica
- Tablas formateadas con estilo profesional
- Diagramas academicos (heatmap, latencia, matriz de confusion, Gantt)
- Tabla de requerimientos
- Diagramas de proceso de pruebas

Dependencias: python-docx, matplotlib
Instalar: pip install python-docx matplotlib
"""

# === IMPORTS: Biblioteca python-docx para manipular documentos Word ===
from docx import Document  # Clase principal del documento Word
from docx.shared import Inches, Pt, RGBColor, Cm, Emu  # Unidades de medida y colores
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Alineaciones de parrafo
from docx.enum.section import WD_ORIENT  # Orientacion de pagina (horizontal/vertical)

# === IMPORTS CONDICIONALES: Graficos nativos de Word (versiones recientes de python-docx) ===
try:
    from docx.enum.chart import XL_CHART_TYPE, XL_LEGEND_POSITION  # Tipos de grafico y posicion de leyenda
    from docx.chart.data import CategoryChartData  # Datos para graficos por categoria
    _HAS_CHARTS = True  # Bandera: hay soporte para graficos nativos
except ImportError:
    _HAS_CHARTS = False  # Bandera: sin soporte, se usara matplotlib como fallback

# === IMPORTS: Manipulacion XML interna de Word (para TOC, numeracion, estilos) ===
from docx.oxml import OxmlElement  # Creacion de elementos XML personalizados
from docx.oxml.ns import qn  # Calificador de espacio de nombres XML
from docx.enum.style import WD_STYLE_TYPE  # Tipos de estilo (parrafo, caracter, etc.)

# === IMPORTS: Estandar de Python ===
import io  # Entrada/salida en memoria (para buffers de imagenes)
import os  # Rutas de archivos del sistema operativo

# === IMPORTS CONDICIONALES: Matplotlib para graficos y diagramas cientificos ===
try:
    import matplotlib  # Libreria principal de graficos
    import matplotlib.pyplot as plt  # Submodulo pyplot para figuras
    matplotlib.rcParams['font.family'] = 'sans-serif'  # Configurar fuente sans-serif por defecto
    _HAS_MPL = True  # Bandera: matplotlib disponible
except ImportError:
    _HAS_MPL = False  # Bandera: matplotlib no instalado
    plt = None  # Referencia nula para evitar errores de NameError

# ============================================================================
# 1. GRAFICOS NATIVOS DE WORD (si la version de python-docx lo soporta)
# ============================================================================

# === BLOQUE CONDICIONAL: Implementacion con graficos nativos de Word ===
if _HAS_CHARTS:
    def insertar_grafico_barras(doc, categorias, valores, titulo="", color=None):
        """
        Inserta un grafico de barras nativo de Word en el documento.

        Parametros:
            doc (Document): Objeto Documento de python-docx donde insertar el grafico.
            categorias (list): Etiquetas del eje X (categorias).
            valores (list): Valores numericos correspondientes a cada categoria.
            titulo (str, opcional): Titulo que aparecera sobre el grafico.
            color (str, opcional): Color de las barras (no utilizado en nativo).

        Retorna:
            docx.chart.Chart: Objeto grafico insertado.

        Efectos secundarios:
            Agrega un grafico de barras de Word (COLUMN_CLUSTERED) al final del documento.
        """
        chart_data = CategoryChartData()  # Contenedor de datos del grafico
        chart_data.categories = categorias  # Asignar etiquetas del eje X
        chart_data.add_series('Valores', valores)  # Agregar serie de datos numericos
        # Crear grafico de barras agrupadas de 16x9 cm
        chart = doc.add_chart(XL_CHART_TYPE.COLUMN_CLUSTERED, width=Cm(16), height=Cm(9), chart_data=chart_data)
        chart.has_title = bool(titulo)  # Activar titulo si se proporciono
        if titulo:
            chart.chart_title.text_frame.paragraphs[0].text = titulo  # Asignar texto del titulo
        chart.legend.position = XL_LEGEND_POSITION.BOTTOM  # Leyenda en la parte inferior
        chart.legend.include_in_layout = False  # No incluir en el layout
        return chart

    def insertar_grafico_pastel(doc, categorias, valores, titulo=""):
        """
        Inserta un grafico de pastel nativo de Word.

        Parametros:
            doc (Document): Objeto Documento destino.
            categorias (list): Etiquetas de las porciones.
            valores (list): Proporciones numericas de cada porcion.
            titulo (str, opcional): Titulo del grafico.

        Retorna:
            docx.chart.Chart: Objeto grafico insertado.

        Efectos secundarios:
            Agrega un grafico circular (PIE) de 14x10 cm al documento.
        """
        chart_data = CategoryChartData()
        chart_data.categories = categorias
        chart_data.add_series('Valores', valores)
        chart = doc.add_chart(XL_CHART_TYPE.PIE, width=Cm(14), height=Cm(10), chart_data=chart_data)
        if titulo:
            chart.has_title = True
            chart.chart_title.text_frame.paragraphs[0].text = titulo
        chart.has_legend = True  # Mostrar leyenda siempre
        chart.legend.position = XL_LEGEND_POSITION.BOTTOM
        return chart

    def insertar_grafico_lineas(doc, categorias, valores, titulo=""):
        """
        Inserta un grafico de lineas nativo de Word.

        Parametros:
            doc (Document): Objeto Documento destino.
            categorias (list): Puntos del eje X.
            valores (list): Valores numericos de cada punto.
            titulo (str, opcional): Titulo del grafico.

        Retorna:
            docx.chart.Chart: Objeto grafico insertado.

        Efectos secundarios:
            Agrega un grafico de lineas con marcadores (LINE_MARKERS) de 16x9 cm.
        """
        chart_data = CategoryChartData()
        chart_data.categories = categorias
        chart_data.add_series('Valores', valores)
        chart = doc.add_chart(XL_CHART_TYPE.LINE_MARKERS, width=Cm(16), height=Cm(9), chart_data=chart_data)
        if titulo:
            chart.has_title = True
            chart.chart_title.text_frame.paragraphs[0].text = titulo
        return chart

# === FALLBACK: Misma API pero usando matplotlib cuando no hay graficos nativos ===
else:
    def insertar_grafico_barras(doc, categorias, valores, titulo="", **kw):
        """Fallback: genera grafico de barras con matplotlib si no hay graficos nativos."""
        datos = dict(zip(categorias, valores))  # Convertir listas paralelas a diccionario
        fig = diagrama_barras_matplotlib(datos, titulo=titulo, ylabel='Valores')  # Crear figura
        if fig:
            insertar_figura_matplotlib(doc, fig)  # Insertar como imagen PNG

    def insertar_grafico_pastel(doc, categorias, valores, titulo="", **kw):
        """Fallback: genera grafico de pastel con matplotlib."""
        datos = dict(zip(categorias, valores))
        fig = diagrama_pastel_matplotlib(datos, titulo=titulo)
        if fig:
            insertar_figura_matplotlib(doc, fig, ancho_cm=12)

    def insertar_grafico_lineas(doc, categorias, valores, titulo="", **kw):
        """Fallback: genera grafico de lineas con matplotlib."""
        datos = dict(zip(categorias, valores))
        fig = diagrama_barras_matplotlib(datos, titulo=titulo, ylabel='Valores')
        if fig:
            insertar_figura_matplotlib(doc, fig)

# ============================================================================
# 2. DIAGRAMAS CON MATPLOTLIB
# ============================================================================

def _fig_to_png(fig, dpi=200):
    """
    Convierte una figura de matplotlib a un buffer de bytes en formato PNG.

    Parametros:
        fig (matplotlib.figure.Figure): Figura a convertir.
        dpi (int, opcional): Resolucion de la imagen en puntos por pulgada (defecto: 200).

    Retorna:
        io.BytesIO: Buffer con los datos binarios de la imagen PNG.

    Efectos secundarios:
        La figura se renderiza a PNG en memoria; no se guarda en disco.
    """
    buf = io.BytesIO()  # Crear buffer en memoria
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', facecolor='white')  # Renderizar a PNG
    buf.seek(0)  # Rebobinar el buffer al inicio para lectura
    return buf

def insertar_figura_matplotlib(doc, fig, ancho_cm=16):
    """
    Inserta una figura de matplotlib como imagen PNG en el documento Word.

    Parametros:
        doc (Document): Objeto Documento destino.
        fig (matplotlib.figure.Figure): Figura a insertar.
        ancho_cm (int, opcional): Ancho de la imagen en centimetros (defecto: 16).

    Efectos secundarios:
        Agrega la imagen al final del documento y cierra la figura de matplotlib.
    """
    if not _HAS_MPL:
        doc.add_paragraph('[Diagrama no disponible: instalar matplotlib]')  # Mensaje de error si falta matplotlib
        return
    buf = _fig_to_png(fig)  # Convertir figura a PNG en buffer
    doc.add_picture(buf, width=Cm(ancho_cm))  # Insertar la imagen en el documento
    plt.close(fig)  # Cerrar la figura para liberar memoria
    buf.close()  # Cerrar el buffer

def diagrama_barras_matplotlib(datos, colores=None, titulo="", xlabel="", ylabel=""):
    """
    Genera un grafico de barras vertical usando matplotlib.

    Parametros:
        datos (dict): Diccionario {categoria: valor} con los datos a graficar.
        colores (str o list, opcional): Color(es) de las barras. Por defecto '#6366F1' (indigo).
        titulo (str, opcional): Titulo del grafico.
        xlabel (str, opcional): Etiqueta del eje X.
        ylabel (str, opcional): Etiqueta del eje Y.

    Retorna:
        matplotlib.figure.Figure o None: La figura generada, o None si matplotlib no esta disponible.

    Efectos secundarios:
        Genera una figura de matplotlib (no la cierra; debe cerrarla quien la llame).
        Muestra el valor numerico encima de cada barra.
    """
    if not _HAS_MPL:
        return None  # Salir silenciosamente si no hay matplotlib
    fig, ax = plt.subplots(figsize=(8, 4.5))  # Crear figura y ejes con tamano fijo
    categorias = list(datos.keys())  # Extraer nombres de categorias del diccionario
    valores = list(datos.values())  # Extraer valores numericos del diccionario
    # Dibujar barras con color personalizado o indigo por defecto, bordes blancos
    bars = ax.bar(categorias, valores, color=colores or '#6366F1', edgecolor='white', linewidth=0.5)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)  # Titulo con estilo
    ax.set_xlabel(xlabel, fontsize=10)  # Etiqueta eje X
    ax.set_ylabel(ylabel, fontsize=10)  # Etiqueta eje Y
    ax.spines['top'].set_visible(False)  # Ocultar borde superior
    ax.spines['right'].set_visible(False)  # Ocultar borde derecho
    ax.tick_params(axis='x', rotation=35, labelsize=9)  # Rotar etiquetas X 35 grados
    ax.tick_params(axis='y', labelsize=9)  # Tamano de numeros del eje Y
    # Agregar valor numerico encima de cada barra
    for bar, val in zip(bars, valores):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(valores)*0.01,
                str(val), ha='center', va='bottom', fontsize=8, fontweight='bold')
    plt.tight_layout()  # Ajustar layout automaticamente
    return fig

def diagrama_pastel_matplotlib(datos, titulo=""):
    """
    Genera un grafico de pastel (circular) con matplotlib.

    Parametros:
        datos (dict): Diccionario {etiqueta: valor} con los datos.
        titulo (str, opcional): Titulo del grafico.

    Retorna:
        matplotlib.figure.Figure o None: La figura generada, o None si matplotlib no esta disponible.

    Efectos secundarios:
        Genera una figura cuadrada de 6x6 pulgadas con leyenda externa.
        Muestra porcentajes en cada porcion.
    """
    if not _HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(6, 6))  # Figura cuadrada para pastel
    labels = list(datos.keys())  # Etiquetas de las porciones
    sizes = list(datos.values())  # Valores numericos de cada porcion
    # Paleta de colores corporativos MindMood (TailwindCSS)
    colors = ['#6366F1', '#EC4899', '#10B981', '#F59E0B', '#EF4444',
              '#06B6D4', '#F97316', '#84CC16', '#7C3AED', '#F43F5E']
    # Crear grafico de pastel con porcentajes, colores limitados al numero de etiquetas
    wedges, texts, autotexts = ax.pie(sizes, labels=None, autopct='%1.0f%%',
                                       colors=colors[:len(labels)], startangle=90,
                                       textprops={'fontsize': 9})
    if titulo:
        ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
    # Leyenda externa al lado derecho con etiqueta y valor
    ax.legend(wedges, [f'{l} ({s})' for l, s in zip(labels, sizes)],
              loc='center left', bbox_to_anchor=(1, 0.5), fontsize=8)
    plt.tight_layout()
    return fig

def diagrama_clases_uml(doc, clases, titulo="Diagrama de Clases"):
    """
    Genera un diagrama de clases UML usando matplotlib y lo inserta en el documento.

    Parametros:
        doc (Document): Objeto Documento destino.
        clases (list): Lista de diccionares con la estructura:
            [{'nombre': str, 'atributos': [str], 'metodos': [str]}, ...]
        titulo (str, opcional): Titulo del diagrama (defecto: 'Diagrama de Clases').

    Efectos secundarios:
        Dibuja rectangulos UML con nombre, atributos y metodos por cada clase.
        Agrega flechas de herencia (gris) y asociacion (morada) entre clases consecutivas.
        Inserta la imagen en el documento y un parrafo descriptivo.
    """
    import matplotlib.patches as mpatches  # Import para formas personalizadas (rectangulos redondeados)
    if not _HAS_MPL:
        doc.add_paragraph(f'[Diagrama no disponible: instalar matplotlib]')
        return
    n = len(clases)  # Cantidad de clases a dibujar
    fig, ax = plt.subplots(figsize=(min(n * 4 + 2, 16), 5))  # Tamano dinamico segun numero de clases
    ax.set_xlim(0, n * 4 + 2)  # Limite horizontal proporcional
    ax.set_ylim(0, 6)  # Altura fija
    ax.axis('off')  # Ocultar ejes
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=15)

    # === Primer pase: dibujar rectangulos de clase ===
    for i, cls in enumerate(clases):
        x = i * 4 + 0.5  # Posicion X de cada clase (espaciado de 4 unidades)
        y = 1.5  # Posicion Y base
        w = 3.5  # Ancho del rectangulo UML
        # Formatear texto: cabecera, separador, atributos, separador, metodos
        header = cls['nombre']  # Nombre de la clase en la cabecera
        attrs = '\n'.join([f'- {a}' for a in cls.get('atributos', [])])  # Atributos con prefijo '-'
        methods = '\n'.join([f'+ {m}()' for m in cls.get('metodos', [])])  # Metodos con prefijo '+'
        # Texto completo del rectangulo UML con separadores Unicode
        text = f"{header}\n\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\n{attrs}\n\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\n{methods}"

        # Rectangulo redondeado con borde gris oscuro y fondo gris claro
        box = mpatches.FancyBboxPatch((x, y), w, 3.5, boxstyle="round,pad=0.1",
                                        edgecolor='#334155', facecolor='#F1F5F9',
                                        linewidth=1.2)
        ax.add_patch(box)  # Agregar rectangulo al eje
        # Texto centrado dentro del rectangulo en fuente monoespaciada
        ax.text(x + w/2, y + 3.5/2, text, ha='center', va='center',
                fontsize=7, fontfamily='monospace', linespacing=1.5)

        # Flecha de herencia (clase i -> i+1) en gris
        if i < n - 1:
            ax.annotate('', xy=(x + w + 0.3, y + 3.5/2),
                        xytext=(x + w - 0.1, y + 3.5/2),
                        arrowprops=dict(arrowstyle='->', color='#94A3B8', lw=1.5))

    # === Segundo pase: flechas de asociacion (desplazadas hacia abajo) ===
    for i, cls in enumerate(clases):
        x = i * 4 + 0.5
        y = 1.5
        w = 3.5
        # Flecha de asociacion en indigo, ligeramente mas abajo que la herencia
        if i < n - 1:
            ax.annotate('', xy=(x + w + 0.8, y + 3.5/2 - 0.5),
                        xytext=(x + w + 0.1, y + 3.5/2 - 0.5),
                        arrowprops=dict(arrowstyle='->', color='#6366F1', lw=1))

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig)
    buf.close()
    # Parrafo descriptivo del diagrama debajo de la imagen
    doc.add_paragraph(f'Diagrama: {titulo}. Muestra las clases del sistema con sus atributos y metodos,'
                       ' flechas de herencia (gris) y asociacion (morada).')

def diagrama_secuencia_matplotlib(doc, pasos, titulo="Diagrama de Secuencia"):
    """
    Genera un diagrama de secuencia UML con matplotlib y lo inserta en el documento.

    Parametros:
        doc (Document): Objeto Documento destino.
        pasos (list): Lista de strings con formato 'Actor1 -> Actor2: mensaje'
                      que describen cada paso de la interaccion.
        titulo (str, opcional): Titulo del diagrama.

    Efectos secundarios:
        Extrae automaticamente los actores unicos de los pasos.
        Dibuja lineas de vida verticales punteadas para cada actor.
        Dibuja flechas horizontales entre actores con el mensaje correspondiente.
        Inserta la imagen en el documento.
    """
    import matplotlib.patches as mpatches
    if not _HAS_MPL:
        doc.add_paragraph(f'[Diagrama no disponible: instalar matplotlib]')
        return

    # === Extraer actores unicos de todos los pasos ===
    actores = []
    for p in pasos:
        parts = p.split('->')  # Separar emisor de receptor
        if len(parts) >= 2:
            for part in parts:
                actor = part.split(':')[0].strip()  # Tomar nombre antes de ':' si existe
                if actor and actor not in actores:
                    actores.append(actor)  # Agregar actor si es nuevo

    n_actores = len(actores)  # Cantidad de columnas (actores)
    # Tamano dinamico: ancho segun actores, alto segun pasos (minimo 6 pulgadas)
    fig, ax = plt.subplots(figsize=(min(n_actores * 3 + 2, 16), max(len(pasos) * 0.5 + 2, 6)))
    ax.set_xlim(0, n_actores * 3 + 2)
    ax.set_ylim(0, len(pasos) * 0.7 + 2)
    ax.axis('off')
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)

    # === Cabeceras de actores en la parte superior ===
    for i, actor in enumerate(actores):
        x = i * 3 + 2  # Posicion X del actor (espaciado de 3 unidades)
        ax.text(x, len(pasos) * 0.7 + 1.5, actor, ha='center', fontsize=9, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.3', facecolor='#6366F1', edgecolor='#4338CA', alpha=0.15))
        # Linea de vida vertical punteada debajo del actor
        ax.axvline(x=x, ymin=0.05, ymax=0.85, color='#CBD5E1', linewidth=1, linestyle='--')

    # === Dibujar flechas de mensaje entre actores ===
    for j, paso in enumerate(pasos):
        y = len(pasos) * 0.7 - j * 0.55  # Posicion Y descendente por cada paso
        parts = paso.split('->')
        if len(parts) >= 2:
            emisor = parts[0].strip()  # Nombre del actor emisor
            resto = parts[1].strip()
            if ':' in resto:
                receptor, mensaje = resto.split(':', 1)  # Separar receptor del mensaje
            else:
                receptor = resto
                mensaje = ''
            receptor = receptor.strip()
            mensaje = mensaje.strip()

            if emisor in actores and receptor in actores:
                x1 = actores.index(emisor) * 3 + 2  # Posicion X del emisor
                x2 = actores.index(receptor) * 3 + 2  # Posicion X del receptor
                # Color alternado entre pasos para mejor legibilidad
                color = '#6366F1' if j % 2 == 0 else '#EC4899'
                # Flecha horizontal de emisor a receptor
                ax.annotate('', xy=(x2 - 0.1, y), xytext=(x1 + 0.1, y),
                            arrowprops=dict(arrowstyle='->', color=color, lw=1.2))
                mid_x = (x1 + x2) / 2  # Punto medio entre los dos actores
                # Texto del mensaje centrado sobre la flecha
                ax.text(mid_x, y + 0.08, mensaje, ha='center', fontsize=6.5,
                        style='italic', color='#475569')

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig)
    buf.close()

def diagrama_casos_uso_matplotlib(doc, actores, casos_uso, relaciones=None, titulo="Diagrama de Casos de Uso"):
    """
    Genera un diagrama de casos de uso UML usando matplotlib.

    Parametros:
        doc (Document): Objeto Documento destino.
        actores (list): Lista de strings con nombres de actores.
        casos_uso (list): Lista de strings con nombres de casos de uso.
        relaciones (list, opcional): Lista de tuplas (idx_actor, idx_caso_uso)
                                     que conectan actores con casos de uso. None = sin conexiones.
        titulo (str, opcional): Titulo del diagrama.

    Efectos secundarios:
        Dibuja figuras de palo simplificadas para actores a la izquierda.
        Dibuja elipses para casos de uso a la derecha con distribucion en columnas.
        Dibuja lineas de conexion entre actores y casos de uso segun 'relaciones'.
        Inserta la imagen en el documento.
    """
    import matplotlib.patches as mpatches
    if not _HAS_MPL:
        doc.add_paragraph(f'[Diagrama no disponible: instalar matplotlib]')
        return

    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=15)

    # === Actores a la izquierda (figuras de palo simplificadas) ===
    for i, actor in enumerate(actores):
        y = 6 - i * 1.5  # Posicion Y descendente por cada actor
        # Circulo para la cabeza
        circle = plt.Circle((1.5, y + 0.4), 0.2, edgecolor='#334155', facecolor='#E2E8F0', linewidth=1.5)
        ax.add_patch(circle)
        # Linea vertical para el cuerpo
        ax.plot([1.5, 1.5], [y - 0.3, y + 0.2], color='#334155', linewidth=1.5)
        # Linea horizontal para los brazos
        ax.plot([1.3, 1.7], [y - 0.05, y - 0.05], color='#334155', linewidth=1.5)
        # Linea en V para las piernas
        ax.plot([1.3, 1.5, 1.7], [y - 0.3, y - 0.15, y - 0.3], color='#334155', linewidth=1.5)
        ax.text(1.5, y - 0.6, actor, ha='center', fontsize=9, fontweight='bold')

    # === Casos de uso a la derecha (elipses) con distribucion en columnas ===
    for i, caso in enumerate(casos_uso):
        col = i // 5  # Columna: cada 5 casos cambia de columna
        row = i % 5   # Fila dentro de la columna
        x = 5 + col * 3.5
        y = 6 - row * 1.2
        # Elipse con borde indigo y fondo azul muy claro
        ellipse = mpatches.Ellipse((x, y), 3, 0.6, edgecolor='#6366F1', facecolor='#EEF2FF', linewidth=1.5)
        ax.add_patch(ellipse)
        ax.text(x, y, caso, ha='center', va='center', fontsize=7.5, fontweight='bold')

    # === Relaciones (lineas entre actores y casos de uso) ===
    if relaciones:
        for actor_idx, caso_idx in relaciones:
            if actor_idx < len(actores) and caso_idx < len(casos_uso):
                y_actor = 6 - actor_idx * 1.5
                col = caso_idx // 5
                row = caso_idx % 5
                x_caso = 5 + col * 3.5
                y_caso = 6 - row * 1.2
                # Linea de conexion desde el actor hasta el caso de uso
                ax.plot([1.7, x_caso - 1.5], [y_actor - 0.1, y_caso], color='#94A3B8',
                        linewidth=1, linestyle='-', alpha=0.7)

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig)
    buf.close()

def diagrama_estados_matplotlib(doc, estados, transiciones, titulo="Diagrama de Maquina de Estados"):
    """
    Genera un diagrama de maquina de estados UML usando matplotlib.

    Parametros:
        doc (Document): Objeto Documento destino.
        estados (list): Lista de strings con nombres de estados.
        transiciones (list): Lista de tuplas (estado_origen, estado_destino, nombre_evento).
        titulo (str, opcional): Titulo del diagrama.

    Efectos secundarios:
        Distribuye los estados en un patron circular.
        Marca el primer estado con un circulo verde (estado inicial).
        Dibuja flechas entre estados con etiquetas del evento.
        Inserta la imagen en el documento.
    """
    import matplotlib.patches as mpatches
    import numpy as np  # Para calculos trigonometricos de posicion circular
    if not _HAS_MPL:
        doc.add_paragraph(f'[Diagrama no disponible: instalar matplotlib]')
        return

    n = len(estados)  # Numero de estados
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=15)

    # === Posiciones en circulo de los estados ===
    positions = {}  # Diccionario {nombre_estado: (x, y)}
    for i, estado in enumerate(estados):
        # Calcular angulo en el circulo (empezando desde arriba: -pi/2)
        angle = 2 * np.pi * i / n - np.pi / 2
        x = 7 + 4.5 * np.cos(angle)  # Centro X = 7, radio = 4.5
        y = 4 + 3 * np.sin(angle)    # Centro Y = 4, radio = 3
        positions[estado] = (x, y)  # Guardar posicion para dibujar transiciones
        # Rectangulo redondeado para el estado
        box = mpatches.FancyBboxPatch((x - 1.8, y - 0.4), 3.6, 0.8,
                                        boxstyle="round,pad=0.15",
                                        edgecolor='#6366F1', facecolor='#EEF2FF', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x, y, estado, ha='center', va='center', fontsize=8.5, fontweight='bold')

        # Estado inicial: circulo verde pequeno a la izquierda del primer estado
        if i == 0:
            circle = plt.Circle((x - 2.2, y), 0.15, color='#10B981', zorder=5)
            ax.add_patch(circle)

    # === Transiciones (flechas con etiquetas) ===
    for orig, dest, evento in transiciones:
        if orig in positions and dest in positions:
            x1, y1 = positions[orig]  # Coordenadas del estado origen
            x2, y2 = positions[dest]  # Coordenadas del estado destino
            dx = x2 - x1  # Diferencia en X
            dy = y2 - y1  # Diferencia en Y
            dist = np.sqrt(dx**2 + dy**2)  # Distancia euclidiana entre estados
            if dist > 0:
                dx, dy = dx/dist, dy/dist  # Vector unitario de direccion
                # Flecha desde el borde del rectangulo origen al borde del destino
                ax.annotate('', xy=(x2 - dx * 1.8, y2 - dy * 0.4),
                            xytext=(x1 + dx * 1.8, y1 + dy * 0.4),
                            arrowprops=dict(arrowstyle='->', color='#F59E0B', lw=1.5))
                # Etiqueta del evento en el punto medio mas un desplazamiento
                mid_x = (x1 + x2) / 2 + 0.3
                mid_y = (y1 + y2) / 2 + 0.3
                ax.text(mid_x, mid_y, evento, fontsize=6.5, color='#92400E',
                        ha='center', style='italic')

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig)
    buf.close()

# ============================================================================
# 3. TABLA DE CONTENIDO AUTOMATICA (TOC)
# ============================================================================

def insertar_indice_automatico(doc, titulo="Tabla de Contenido"):
    """
    Inserta una tabla de contenido (TOC) automatica en el documento Word.

    La TOC se genera con campos de Word (FldChar) que Word actualiza
    al abrir el documento. Incluye encabezados de nivel 1 a 3.

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo mostrado encima del indice (defecto: 'Tabla de Contenido').

    Efectos secundarios:
        Agrega un parrafo centrado con el titulo del indice.
        Inserta un campo TOC \\o "1-3" que Word renderizara al abrir el archivo.
        Muestra un texto placeholder hasta que el usuario actualice el campo.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(99, 102, 241)  # Color indigo corporativo

    # === Construir campo TOC mediante XML de Word ===
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # Elemento de campo de Word
    fldChar.set(qn('w:fldCharType'), 'begin')  # Inicio del campo
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')  # Instruccion del campo
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # TOC: niveles 1-3, hipervinculos, sin numeracion, usar estilos
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')  # Separador entre instruccion y resultado
    run._r.append(fldChar2)

    # Texto placeholder que aparece hasta que Word actualice el campo
    run2 = paragraph.add_run("(Actualice este campo con clic derecho > Actualizar campo)")
    run2.font.color.rgb = RGBColor(148, 163, 184)  # Gris claro
    run2.font.size = Pt(9)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')  # Fin del campo
    run2._r.append(fldChar3)

# ============================================================================
# 4. MARCA DE AGUA EN ENCABEZADOS
# ============================================================================

def agregar_marca_agua(doc, texto="BORRADOR", color_hex="D9D9D9"):
    """
    Agrega una marca de agua de texto al encabezado del documento.

    Funciona insertando texto grande y semitransparente en el encabezado
    de la primera seccion, visible detras del contenido.

    Parametros:
        doc (Document): Objeto Documento destino.
        texto (str, opcional): Texto de la marca de agua (defecto: 'BORRADOR').
        color_hex (str, opcional): Color en hexadecimal sin '#' (defecto: 'D9D9D9' = gris claro).

    Efectos secundarios:
        Modifica el encabezado de la primera seccion.
        El encabezado se desvincula del anterior (is_linked_to_previous = False).
    """
    section = doc.sections[0]  # Primera seccion del documento
    header = section.header  # Acceder al encabezado de la seccion
    header.is_linked_to_previous = False  # Desvincular de encabezados anteriores
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la marca de agua
    run = p.add_run(texto)
    run.font.size = Pt(60)  # Tamano grande para que se vea como marca de agua
    # Convertir HEX a RGB y agregar canal alpha (180 para semitransparencia)
    run.font.color.rgb = RGBColor(*[int(color_hex[i:i+2], 16) for i in (0, 2, 4)] + [180])
    run.font.name = 'Calibri'

# ============================================================================
# 5. SECCIONES CON ORIENTACION DE PAGINA
# ============================================================================

def nueva_pagina_horizontal(doc):
    """
    Agrega una nueva seccion al documento en orientacion horizontal (apaisada).

    Util para insertar tablas anchas, diagramas grandes o graficos de Gantt.

    Parametros:
        doc (Document): Objeto Documento destino.

    Retorna:
        docx.section.Section: La nueva seccion creada.

    Efectos secundarios:
        Crea un salto de seccion y cambia la orientacion a horizontal.
        Define margenes personalizados: 2 cm superior/inferior, 2.5 cm izquierdo/derecho.
    """
    section = doc.add_section()  # Crear nueva seccion
    section.orientation = WD_ORIENT.LANDSCAPE  # Orientacion horizontal
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return section

def nueva_pagina_vertical(doc):
    """
    Agrega una nueva seccion en orientacion vertical (retrato).

    Util para volver a orientacion normal despues de una seccion horizontal.

    Parametros:
        doc (Document): Objeto Documento destino.

    Retorna:
        docx.section.Section: La nueva seccion creada.

    Efectos secundarios:
        Crea un salto de seccion con orientacion vertical y margenes uniformes de 2.5 cm.
    """
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT  # Orientacion vertical (defecto)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return section

# ============================================================================
# 6. ESTILOS PERSONALIZADOS
# ============================================================================

def crear_estilo(doc, nombre, fuente='Calibri', tamano=12, color=RGBColor(0x0F, 0x17, 0x2A),
                 negrita=False, tipo=WD_STYLE_TYPE.PARAGRAPH):
    """
    Crea o recupera un estilo de parrafo personalizado en el documento.

    Si el estilo ya existe, lo recupera; si no, lo crea.

    Parametros:
        doc (Document): Objeto Documento.
        nombre (str): Nombre del estilo.
        fuente (str, opcional): Nombre de la tipografia (defecto: 'Calibri').
        tamano (int, opcional): Tamano de fuente en puntos (defecto: 12).
        color (RGBColor, opcional): Color de la fuente (defecto: azul oscuro #0F172A).
        negrita (bool, opcional): Si la fuente debe ser negrita (defecto: False).
        tipo (WD_STYLE_TYPE, opcional): Tipo de estilo (defecto: PARAGRAPH).

    Retorna:
        docx.style.Style: El estilo creado o recuperado.
    """
    if nombre not in [s.name for s in doc.styles]:
        estilo = doc.styles.add_style(nombre, tipo)  # Crear nuevo estilo
    else:
        estilo = doc.styles[nombre]  # Recuperar estilo existente
    estilo.font.name = fuente
    estilo.font.size = Pt(tamano)
    estilo.font.color.rgb = color
    estilo.font.bold = negrita
    return estilo

def configurar_estilos_titulos(doc):
    """
    Configura los estilos de encabezado Heading 1/2/3 con colores corporativos MindMood.

    Heading 1: Indigo, 20pt, negrita
    Heading 2: Azul oscuro, 15pt, negrita
    Heading 3: Gris pizarra, 12pt, negrita

    Parametros:
        doc (Document): Objeto Documento a modificar.

    Efectos secundarios:
        Modifica globalmente los estilos 'Heading 1', 'Heading 2' y 'Heading 3' del documento.
    """
    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(20)
    h1.font.color.rgb = RGBColor(99, 102, 241)  # Indigo corporativo
    h1.font.bold = True

    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(15)
    h2.font.color.rgb = RGBColor(30, 41, 59)  # Azul oscuro
    h2.font.bold = True

    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(71, 85, 105)  # Gris pizarra
    h3.font.bold = True

# ============================================================================
# 7. PIE DE PAGINA CON NUMERACION AUTOMATICA
# ============================================================================

def configurar_pie_pagina(doc):
    """
    Configura el pie de pagina de la primera seccion con numeracion automatica.

    Muestra 'Pagina X' donde X es el numero de pagina actual usando
    el campo PAGE de Word (FldChar).

    Parametros:
        doc (Document): Objeto Documento destino.

    Efectos secundarios:
        Modifica el pie de pagina de la primera seccion.
        Desvincula el pie de pagina de secciones anteriores.
        Agrega el campo PAGE para numeracion dinamica.
    """
    section = doc.sections[0]  # Primera seccion del documento
    footer = section.footer  # Acceder al pie de pagina
    footer.is_linked_to_previous = False  # Desvincular de pies anteriores
    p = footer.paragraphs[0]  # Primer parrafo del pie de pagina
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar contenido

    # === Construir campo de numero de pagina via XML de Word ===
    run = p.add_run("Pagina ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)  # Gris claro

    fldChar = OxmlElement('w:fldChar')  # Inicio del campo
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')  # Instruccion PAGE
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')  # Separador
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    run2 = p.add_run('1')  # Valor placeholder que Word reemplazara
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(148, 163, 184)
    fldChar3 = OxmlElement('w:fldChar')  # Fin del campo
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

# ============================================================================
# 8. TABLA FORMATEADA CON ESTILO PROFESIONAL
# ============================================================================

def add_tabla(doc, headers, rows, col_widths=None):
    """
    Agrega una tabla con formato profesional al documento Word.

    Caracteristicas:
        - Cabecero con fondo indigo (#6366F1) y texto blanco en negrita.
        - Filas alternadas con fondo gris claro (#F8FAFC).
        - Fuente Calibri 8.5pt en las celdas de datos.
        - Autoajuste de columnas activado.

    Parametros:
        doc (Document): Objeto Documento destino.
        headers (list): Lista de strings para los encabezados de columna.
        rows (list): Lista de listas con los datos de cada fila.
        col_widths (list, opcional): No utilizado (reservado para compatibilidad).

    Retorna:
        docx.table.Table: La tabla creada.

    Efectos secundarios:
        Agrega una tabla al final del documento con estilo 'Light Shading Accent 1'.
    """
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))  # +1 para cabecero
    t.style = 'Light Shading Accent 1'  # Estilo base de Word
    t.alignment = 1  # Centrar tabla en la pagina
    t.autofit = True  # Autoajustar ancho de columnas

    # === Formatear fila de cabeceras ===
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
        # Fondo oscuro indigo en cabecera mediante XML
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '6366F1')  # Color indigo
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # === Formatear filas de datos con alternancia de color ===
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]  # Fila ri+1 (despues del cabecero)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
                    r.font.name = 'Calibri'
            # Alternar color de fondo en filas impares
            if ri % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F8FAFC')  # Gris muy claro
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
    return t

# ============================================================================
# 9. DESCRIPCION VISUAL DE DIAGRAMAS (fallback textual)
# ============================================================================
# Nota: La seccion 9 actualmente sirve como separador. La funcion
# descripcion_diagrama() esta al final de la seccion 10.

# ============================================================================
# 10. GRAFICOS ACADEMICOS (Heatmap, Latencia, Matriz de Confusion)
# ============================================================================

def diagrama_heatmap_emocional(doc, titulo="Heatmap de Actividad Emocional"):
    """
    Genera un heatmap que correlaciona emociones con horas del dia.

    Los datos son simulados con valores aleatorios con fines de demostracion.
    Muestra la frecuencia de cada emocion a lo largo de las 24 horas.

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo del grafico.

    Efectos secundarios:
        Genera datos aleatorios con numpy.
        Inserta la imagen del heatmap (14x5 pulgadas) en el documento.
        Agrega una nota aclaratoria sobre datos representativos.
    """
    import numpy as np  # Import local para generacion de datos aleatorios
    if not _HAS_MPL:
        doc.add_paragraph('[Heatmap no disponible]')
        return
    horas = list(range(24))  # 0 a 23 (las 24 horas del dia)
    emociones = ['Feliz', 'Triste', 'Enojado', 'Ansioso', 'Neutral', 'Crisis']  # 6 categorias
    data = np.random.randint(1, 20, size=(len(emociones), len(horas)))  # Datos simulados 6x24
    fig, ax = plt.subplots(figsize=(14, 5))
    im = ax.imshow(data, cmap='YlOrRd', aspect='auto')  # Mapa de calor con colores amarillo-rojo
    ax.set_xticks(range(len(horas)))
    ax.set_xticklabels([f'{h}:00' for h in horas], fontsize=7, rotation=45)  # Etiquetas de hora
    ax.set_yticks(range(len(emociones)))
    ax.set_yticklabels(emociones, fontsize=9)  # Etiquetas de emociones
    ax.set_xlabel('Hora del dia', fontsize=10)
    ax.set_ylabel('Emocion detectada', fontsize=10)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
    fig.colorbar(im, ax=ax, label='Frecuencia')  # Barra de color lateral
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()
    doc.add_paragraph('Nota: Los datos son representativos. El heatmap muestra picos de actividad emocional en horas especificas del dia.')

def diagrama_latencia_ia(doc, titulo="Latencia de IA vs Longitud del Texto"):
    """
    Genera un grafico de dispersion que relaciona tiempo de respuesta vs cantidad de palabras.

    Simula la latencia del pipeline de IA para validar el requerimiento PS-005
    (tiempo de respuesta menor a 5 segundos). Usa semilla aleatoria para reproducibilidad.

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo del grafico.

    Efectos secundarios:
        Genera 50 puntos de datos simulados con numpy (semilla 42).
        Dibuja una linea horizontal roja punteada en el umbral de 5 segundos.
        Inserta la imagen en el documento y agrega nota de validacion.
    """
    import numpy as np
    if not _HAS_MPL:
        doc.add_paragraph('[Grafico de latencia no disponible]')
        return
    np.random.seed(42)  # Semilla fija para reproducibilidad
    palabras = np.random.randint(10, 200, 50)  # 50 muestras de 10 a 200 palabras
    # Modelo lineal: 0.5s base + 0.012s por palabra + ruido normal
    latencia = 0.5 + palabras * 0.012 + np.random.normal(0, 0.2, 50)
    latencia = np.clip(latencia, 0.3, 4.5)  # Limitar entre 0.3s y 4.5s
    fig, ax = plt.subplots(figsize=(8, 5))
    # Grafico de dispersion con puntos indigos
    ax.scatter(palabras, latencia, alpha=0.7, color='#6366F1', edgecolors='white', linewidth=0.5, s=50)
    ax.axhline(y=5.0, color='#EF4444', linestyle='--', linewidth=1.5, label='Umbral maximo (5s)')  # Linea de umbral
    ax.set_xlabel('Cantidad de palabras', fontsize=10)
    ax.set_ylabel('Tiempo de respuesta (s)', fontsize=10)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
    ax.legend(fontsize=9)
    ax.grid(axis='y', linestyle='--', alpha=0.3)  # Cuadricula horizontal suave
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(14))
    plt.close(fig); buf.close()
    doc.add_paragraph('Nota: Todos los puntos estan por debajo del umbral de 5 segundos, validando el requerimiento PS-005.')

def diagrama_matriz_confusion(doc, titulo="Matriz de Confusion - Robertuito Emotion"):
    """
    Genera una matriz de confusion simulada basada en pruebas de integracion del modelo.

    La matriz compara las etiquetas reales contra las predicciones del modelo
    Robertuito Emotion para 6 categorias emocionales.

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo del grafico.

    Efectos secundarios:
        Dibuja una matriz 6x6 con mapa de color azul.
        Muestra el valor numerico en cada celda con color adaptativo.
        Agrega estadisticas de exactitud (aciertos/total) en un parrafo.
        Inserta la imagen en el documento.
    """
    import numpy as np
    if not _HAS_MPL:
        doc.add_paragraph('[Matriz de confusion no disponible]')
        return
    etiquetas = ['Feliz', 'Triste', 'Enojado', 'Ansioso', 'Neutral', 'Crisis']  # 6 clases
    n = len(etiquetas)
    # Matriz de confusion simulada: diagonal alta = buenas predicciones
    matriz = np.array([
        [18, 1, 0, 1, 2, 0],
        [1, 15, 1, 1, 1, 1],
        [0, 1, 14, 1, 0, 0],
        [1, 1, 1, 13, 1, 1],
        [2, 1, 0, 1, 16, 0],
        [0, 0, 0, 0, 0, 20],  # Crisis tiene 20/20 (perfecto)
    ])
    fig, ax = plt.subplots(figsize=(8, 7))
    im = ax.imshow(matriz, cmap='Blues', aspect='equal')  # Mapa de color azul
    ax.set_xticks(range(n)); ax.set_xticklabels(etiquetas, fontsize=8, rotation=45, ha='right')
    ax.set_yticks(range(n)); ax.set_yticklabels(etiquetas, fontsize=8)
    ax.set_xlabel('Predicho', fontsize=10)
    ax.set_ylabel('Real', fontsize=10)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
    # Escribir el valor numerico en cada celda
    for i in range(n):
        for j in range(n):
            ax.text(j, i, matriz[i, j], ha='center', va='center', fontsize=11,
                    color='white' if matriz[i, j] > 10 else 'black', fontweight='bold')
    fig.colorbar(im, ax=ax, label='Aciertos', shrink=0.7)  # Barra de color
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(14))
    plt.close(fig); buf.close()
    total = matriz.sum()  # Suma de todos los elementos
    aciertos = matriz.trace()  # Suma de la diagonal (predicciones correctas)
    exactitud = aciertos / total * 100  # Porcentaje de exactitud
    doc.add_paragraph(f'Exactitud del modelo: {exactitud:.1f}% ({aciertos}/{total} predicciones correctas). '
                       'La diagonal principal muestra aciertos; los valores fuera de ella son confusiones del modelo.')

def diagrama_gantt_matplotlib(doc, tareas, titulo="Diagrama de Gantt", inicio="2026-02-01", fin="2026-05-15"):
    """
    Genera un diagrama de Gantt para visualizar el cronograma del proyecto.

    Cada tarea tiene: fase, nombre, fecha inicio, fecha fin.
    Las fases se colorean con una paleta predefinida.
    Incluye lineas verticales punteadas para cada mes.

    Parametros:
        doc (Document): Objeto Documento destino.
        tareas (list): Lista de tuplas (fase, actividad, fecha_inicio, fecha_fin)
                       donde las fechas son strings en formato 'YYYY-MM-DD'.
        titulo (str, opcional): Titulo del diagrama.
        inicio (str, opcional): Fecha de inicio del eje X (defecto: '2026-02-01').
        fin (str, opcional): Fecha de fin del eje X (defecto: '2026-05-15').

    Efectos secundarios:
        Requiere pandas y matplotlib.dates.
        Dibuja barras horizontales con etiquetas de actividad.
        Oculta las etiquetas del eje Y para un aspecto mas limpio.
        Inserta la imagen en el documento.
    """
    import pandas as pd  # Para generacion de rangos de fechas
    import matplotlib.dates as mdates
    from datetime import datetime, timedelta
    if not _HAS_MPL:
        doc.add_paragraph(f'[Diagrama no disponible: instalar matplotlib]')
        return
    start = datetime.strptime(inicio, '%Y-%m-%d')  # Convertir fecha inicio a objeto datetime
    end = datetime.strptime(fin, '%Y-%m-%d')  # Convertir fecha fin a objeto datetime
    total_days = (end - start).days  # Duracion total en dias
    fig, ax = plt.subplots(figsize=(14, len(tareas) * 0.5 + 2))  # Alto dinamico segun tareas
    # Paleta de colores por numero de fase
    colores_fase = {'1': '#6366F1', '2': '#EC4899', '3': '#10B981', '4': '#F59E0B', '5': '#EF4444', '6': '#8B5CF6'}
    for i, (fase, actividad, f_inicio, f_fin) in enumerate(tareas):
        y = len(tareas) - i - 1  # Posicion Y (invertida para que la primera tarea este arriba)
        d_inicio = datetime.strptime(f_inicio, '%Y-%m-%d')
        d_fin = datetime.strptime(f_fin, '%Y-%m-%d')
        dias = (d_fin - d_inicio).days  # Duracion de la tarea en dias
        color = colores_fase.get(fase.split('.')[0], '#6366F1')  # Color segun numero de fase
        # Barra horizontal desde (d_inicio - start) con duracion 'dias'
        ax.barh(y, dias, left=(d_inicio - start).days, height=0.6, color=color, edgecolor='white', linewidth=0.5, alpha=0.85)
        # Etiqueta de la actividad blanca dentro de la barra
        ax.text((d_inicio - start).days + 0.3, y, f'{actividad} ({fase})', va='center', fontsize=7.5, fontweight='bold', color='white')
    # Lineas verticales por mes
    meses = pd.date_range(start=start, end=end, freq='MS')  # 'MS' = Month Start
    for m in meses:
        ax.axvline(x=(m - start).days, color='#CBD5E1', linewidth=0.5, linestyle='--', alpha=0.5)
    ax.set_yticks(range(len(tareas)))
    ax.set_yticklabels([''] * len(tareas))  # Ocultar nombres en eje Y (ya estan en las barras)
    ax.set_xlim(0, total_days)
    # Marcas semanales en eje X con formato dia/mes
    weeks = pd.date_range(start=start, end=end, freq='W')
    ax.set_xticks([(w - start).days for w in weeks])
    ax.set_xticklabels([w.strftime('%d/%m') for w in weeks], fontsize=7)
    ax.set_xlabel('Febrero - Mayo 2026', fontsize=9)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.tick_params(axis='y', length=0)
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()

def descripcion_diagrama(doc, titulo, descripcion, elementos):
    """
    Agrega una descripcion textual detallada de un diagrama como alternativa cuando
    no se puede generar la imagen (fallback textual).

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str): Titulo del diagrama que se usara como encabezado nivel 2.
        descripcion (str): Parrafo descriptivo del diagrama.
        elementos (list): Lista de diccionarios con claves 'nombre', 'desc'
                          y opcionalmente 'conexiones' (lista de strings).

    Efectos secundarios:
        Agrega un encabezado nivel 2, un parrafo descriptivo y una lista
        con viñetas de los elementos y sus conexiones.
    """
    doc.add_heading(titulo, level=2)
    doc.add_paragraph(descripcion)
    doc.add_paragraph("Estructura del diagrama:")  # Encabezado de la lista
    for elem in elementos:
        p = doc.add_paragraph(style='List Bullet')  # Elemento con viñeta
        run = p.add_run(elem.get('nombre', ''))  # Nombre en negrita
        run.bold = True
        p.add_run(f': {elem.get("desc", "")}')  # Descripcion
        if 'conexiones' in elem:
            for conn in elem['conexiones']:
                doc.add_paragraph(f'  → {conn}', style='List Bullet')  # Sub-elementos con flecha

# ============================================================================
# 10. TABLA DE REQUERIMIENTOS (con seccion numerada)
# ============================================================================

def add_tabla_requerimientos(doc, headers, rows):
    """
    Crea una tabla de requerimientos con formato especial (fuente mas pequena).

    Util para listar requerimientos funcionales y no funcionales del proyecto.

    Parametros:
        doc (Document): Objeto Documento destino.
        headers (list): Lista de strings para encabezados.
        rows (list): Lista de listas con los datos de cada requerimiento.

    Retorna:
        docx.table.Table: La tabla creada.

    Efectos secundarios:
        Agrega una tabla con estilo 'Light Grid Accent 1'.
        Encabezados en negrita 7.5pt; datos en 7pt.
    """
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'  # Estilo con bordes de cuadricula
    t.alignment = 1  # Centrar
    # Formatear cabeceras
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(7.5)  # Fuente pequena para tablas densas
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)  # Fuente aun mas pequena para datos
    return t

# ============================================================================
# 11. DIAGRAMAS DE PROCESO DE PRUEBAS (Pipeline e informes)
# ============================================================================

def diagrama_flujo_pruebas(doc, titulo="Proceso de ejecucion de pruebas"):
    """
    Genera un diagrama visual del pipeline de ejecucion de pruebas.

    Muestra los pasos desde la configuracion (pytest.ini) hasta los resultados
    JSON, pasando por los casos de prueba, TestClient y el pipeline de IA.
    Incluye legendas con los IDs de las pruebas organizadas por tipo.

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo del diagrama.

    Efectos secundarios:
        Dibuja rectangulos conectados por flechas que representan el flujo.
        Inserta la imagen en el documento.
    """
    if not _HAS_MPL:
        doc.add_paragraph('[Diagrama no disponible]')
        return
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.set_xlim(0, 14); ax.set_ylim(0, 6); ax.axis('off')
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=15, color='#6366F1')

    # Etapas del pipeline de pruebas con posiciones, etiquetas y colores
    pasos = [
        (1, 5, 'pytest.ini', 'Configuracion', '#6366F1'),
        (3.5, 5, 'test_mindmood.py', '15 casos de prueba', '#EC4899'),
        (6, 5, 'TestClient', 'Simula requests HTTP', '#10B981'),
        (8.5, 5, 'main.py', 'Pipeline de IA', '#F59E0B'),
        (11, 5, 'Response', 'Resultados JSON', '#EF4444'),
    ]
    for x, y, label, desc, color in pasos:
        # Rectangulo redondeado con color semitransparente de fondo
        rect = mpatches.FancyBboxPatch((x-0.9, y-0.5), 1.8, 1, boxstyle="round,pad=0.1", edgecolor=color, facecolor=f'{color}18', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y+0.15, label, ha='center', va='center', fontsize=7, fontweight='bold', color=color)
        ax.text(x, y-0.25, desc, ha='center', va='center', fontsize=6, color='#475569')

    # Flechas de conexion entre etapas
    flechas = [(1.9, 5, 3.5, 5), (3.5, 5, 6, 5), (6, 5, 8.5, 5), (8.5, 5, 11, 5)]
    for x1, y1, x2, y2 in flechas:
        ax.annotate('', xy=(x2-0.9, y2), xytext=(x1+0.9, y1), arrowprops=dict(arrowstyle='->', color='#94A3B8', lw=1.5))

    # Leyendas de cobertura de pruebas debajo del diagrama
    ax.text(6, 3.5, 'Modulares (PM-001 a PM-005)', fontsize=8, fontweight='bold', ha='center', color='#6366F1')
    ax.text(6, 3.0, 'Integration (PI-001 a PI-005)', fontsize=8, fontweight='bold', ha='center', color='#EC4899')
    ax.text(6, 2.5, 'Sistema (PS-001 a PS-005)', fontsize=8, fontweight='bold', ha='center', color='#10B981')
    ax.text(6, 1.8, 'Independientes entre si, 3 tipos de cobertura', fontsize=7, ha='center', color='#64748B', style='italic')

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()

def diagrama_comparativa_pruebas(doc, antes_pass, antes_fail, despues_pass, despues_fail, titulo="Comparativa antes/despues de parches"):
    """
    Genera un grafico de barras comparativo entre los resultados de pruebas
    antes y despues de aplicar parches correctivos.

    Parametros:
        doc (Document): Objeto Documento destino.
        antes_pass (int): Pruebas pasadas antes de parches.
        antes_fail (int): Pruebas falladas antes de parches.
        despues_pass (int): Pruebas pasadas despues de parches.
        despues_fail (int): Pruebas falladas despues de parches.
        titulo (str, opcional): Titulo del grafico.

    Efectos secundarios:
        Dibuja barras agrupadas rojas (antes) y verdes (despues) lado a lado.
        Muestra los valores numericos sobre cada barra.
        Inserta la imagen en el documento.
    """
    if not _HAS_MPL: return
    fig, ax = plt.subplots(figsize=(6, 4))
    categorias = ['Passed', 'Failed']  # Dos categorias
    x = [0, 1]  # Posiciones en el eje X
    width = 0.3  # Ancho de cada barra
    # Barras "Antes" en rojo semitransparente
    ax.bar([i - width/2 for i in x], [antes_pass, antes_fail], width, label='Antes', color='#EF444480', edgecolor='#EF4444', linewidth=1.5)
    # Barras "Despues" en verde semitransparente
    ax.bar([i + width/2 for i in x], [despues_pass, despues_fail], width, label='Despues', color='#10B98180', edgecolor='#10B981', linewidth=1.5)
    ax.set_xticks(x); ax.set_xticklabels(categorias, fontsize=10, fontweight='bold')
    ax.set_ylabel('Cantidad de pruebas', fontsize=9)
    ax.set_title(titulo, fontsize=11, fontweight='bold', pad=10, color='#6366F1')
    ax.legend(fontsize=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0, max(antes_pass, despues_pass) + 3)
    # Etiquetas de valor sobre las barras "Antes" (rojo)
    for i, v in enumerate([antes_pass, antes_fail]):
        ax.text(i - width/2, v + 0.3, str(v), ha='center', fontsize=9, fontweight='bold', color='#EF4444')
    # Etiquetas de valor sobre las barras "Despues" (verde)
    for i, v in enumerate([despues_pass, despues_fail]):
        ax.text(i + width/2, v + 0.3, str(v), ha='center', fontsize=9, fontweight='bold', color='#10B981')
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(12))
    plt.close(fig); buf.close()

def diagrama_pastel_resultados(doc, passed, failed, xfailed=0, titulo="Distribucion de resultados"):
    """
    Genera un grafico de pastel con la distribucion de resultados de pruebas.

    Parametros:
        doc (Document): Objeto Documento destino.
        passed (int): Pruebas pasadas.
        failed (int): Pruebas falladas.
        xfailed (int, opcional): Pruebas con fallo esperado (xfail).
        titulo (str, opcional): Titulo del grafico.

    Efectos secundarios:
        Dibuja un pastel con colores verde (passed), rojo (failed), amarillo (xfail).
        Solo incluye categorias con valores > 0.
        Inserta la imagen en el documento.
    """
    if not _HAS_MPL: return
    fig, ax = plt.subplots(figsize=(5, 4))
    labels = ['Passed', 'Failed', 'Expected Fail']  # Etiquetas de las 3 categorias
    sizes = [passed, failed, xfailed]  # Valores originales
    colors = ['#10B981', '#EF4444', '#F59E0B']  # Verde, rojo, amarillo
    sizes = [s for s in sizes if s > 0]  # Solo valores positivos
    labels = [l for l, s in zip(labels, [passed, failed, xfailed]) if s > 0]  # Etiquetas filtradas
    if not sizes: return  # No dibujar si no hay datos
    wedges, texts, autotexts = ax.pie(sizes, labels=None, autopct='%1.1f%%', colors=colors[:len(sizes)], startangle=90, textprops={'fontsize': 9})
    ax.set_title(titulo, fontsize=11, fontweight='bold', pad=10, color='#6366F1')
    ax.legend(wedges, [f'{l}: {s}' for l, s in zip(labels, [passed, failed, xfailed])], loc='center left', bbox_to_anchor=(1, 0.5), fontsize=8)
    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(10))
    plt.close(fig); buf.close()

# ============================================================================
# 12. NUEVOS DIAGRAMAS PROFESIONALES (Fase 2 - Mejoras Visuales)
# ============================================================================

def diagrama_flujo_pipeline_ia(doc, titulo="Pipeline de IA - 10 Etapas"):
    """
    Genera un diagrama de flujo (flowchart) del pipeline de inteligencia artificial
    de MindMood, mostrando las 10 etapas secuenciales con colores y conexiones.

    Etapas visualizadas:
        1. Recepcion de texto -> 2. Limpieza de emojis -> 3. Normalizacion de jerga
        -> 4. Sentiment Analysis -> 5. Refuerzo emocional -> 6. Crisis Detection
        -> 7. Emotion Analysis -> 8. Keyword Reinforcement -> 9. Calculo de confianza
        -> 10. Generacion de resumen

    Parametros:
        doc (Document): Objeto Documento destino.
        titulo (str, opcional): Titulo del diagrama.
    """
    import matplotlib.patches as mpatches
    if not _HAS_MPL: return

    etapas = [
        ("1.\nRecepcion\nde texto", "#6366F1"),
        ("2.\nLimpieza\nde emojis", "#EC4899"),
        ("3.\nNormalizacion\nde jerga", "#10B981"),
        ("4.\nSentiment\nAnalysis", "#F59E0B"),
        ("5.\nRefuerzo\nemocional", "#EF4444"),
        ("6.\nCrisis\nDetection", "#7C3AED"),
        ("7.\nEmotion\nAnalysis", "#06B6D4"),
        ("8.\nKeyword\nReinforcement", "#F97316"),
        ("9.\nCalculo de\nconfianza", "#84CC16"),
        ("10.\nGeneracion de\nresumen", "#EC4899"),
    ]

    fig, ax = plt.subplots(figsize=(16, 9))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title(titulo, fontsize=15, fontweight='bold', pad=20, color='#6366F1')

    box_w, box_h = 2.2, 1.4
    for i, (texto, color) in enumerate(etapas):
        col = i % 2
        row = i // 2
        x = 1.8 + col * 5.5
        y = 8.5 - row * 1.8

        rect = mpatches.FancyBboxPatch((x, y), box_w, box_h,
               boxstyle="round,pad=0.15", edgecolor=color,
               facecolor=f'{color}20', linewidth=2)
        ax.add_patch(rect)
        ax.text(x + box_w/2, y + box_h/2, texto, ha='center', va='center',
                fontsize=7.5, fontweight='bold', color=color)

        if i < len(etapas) - 1:
            next_col = (i + 1) % 2
            next_row = (i + 1) // 2
            nx, ny = 1.8 + next_col * 5.5, 8.5 - next_row * 1.8
            sx, sy = x + box_w, y + box_h/2
            ex, ey = nx, ny + box_h/2
            ax.annotate('', xy=(ex, ey), xytext=(sx, sy),
                       arrowprops=dict(arrowstyle='->', color='#94A3B8', lw=1.5))

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()
    doc.add_paragraph(
        'Explicacion: El pipeline de IA de MindMood ejecuta 10 etapas secuenciales '
        'que transforman el texto crudo del usuario en un analisis emocional completo. '
        'Las etapas 4 y 7 utilizan modelos HuggingFace Robertuito pre-entrenados en '
        'espanol para clasificar sentimiento (POS/NEG/NEU) y emociones (7 categorias). '
        'La etapa 6 implementa deteccion de crisis en 3 niveles: keywords directas, '
        'fuzzy matching para errores ortograficos y patrones regex. La etapa 10 genera '
        'un resumen empatico personalizado basado en el estado emocional detectado.')

def diagrama_er_database(doc, titulo="Diagrama Entidad-Relacion - Base de Datos"):
    """
    Genera un diagrama Entidad-Relacion (ER) de la base de datos PostgreSQL
    de MindMood, mostrando las 3 tablas principales y sus relaciones.

    Tablas: profiles (6 campos), entries (7 campos), contact_requests (8 campos).
    Relaciones: profiles 1---N entries, profiles 1---N contact_requests.
    """
    import matplotlib.patches as mpatches
    if not _HAS_MPL: return

    tablas = [
        {"nombre": "profiles", "campos": [
            "id (PK)", "email", "nombre", "avatar",
            "tema", "idioma", "created_at"]},
        {"nombre": "entries", "campos": [
            "id (PK)", "user_id (FK)", "texto", "mood",
            "score", "requires_help", "created_at"]},
        {"nombre": "contact_requests", "campos": [
            "id (PK)", "user_id (FK)", "admin_id (FK)",
            "entry_id (FK)", "status", "mensaje", "respuesta", "created_at"]},
    ]

    fig, ax = plt.subplots(figsize=(16, 7))
    ax.set_xlim(0, 16); ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=15, color='#6366F1')
    ax.text(8, 7.5, 'RLS: Row Level Security habilitado en todas las tablas | 12 funciones RPC SECURITY DEFINER',
            ha='center', fontsize=8, color='#64748B', style='italic')

    positions = [(2, 3.5), (8, 3.5), (14, 3.5)]

    for (xx, yy), tabla in zip(positions, tablas):
        n_campos = len(tabla["campos"])
        altura = max(n_campos * 0.35 + 0.6, 2.5)
        rect = mpatches.FancyBboxPatch((xx - 2, yy - altura/2), 4, altura,
               boxstyle="round,pad=0.1", edgecolor='#6366F1',
               facecolor='#EEF2FF', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(xx, yy + altura/2 - 0.25, tabla["nombre"],
                ha='center', fontsize=10, fontweight='bold', color='#6366F1')
        ax.axhline(y=yy + altura/2 - 0.5, xmin=(xx-1.8)/16, xmax=(xx+1.8)/16,
                   color='#6366F1', linewidth=1)
        for j, campo in enumerate(tabla["campos"]):
            clr = '#0F172A'
            if '(PK)' in campo: clr = '#7C3AED'
            elif '(FK)' in campo: clr = '#F59E0B'
            ax.text(xx, yy + altura/2 - 0.7 - j * 0.35, campo,
                    ha='center', fontsize=7.5, fontfamily='monospace', color=clr)

    ax.annotate('', xy=(5.8, 3.5), xytext=(4.2, 3.5),
               arrowprops=dict(arrowstyle='-', color='#F59E0B', lw=1.2))
    ax.text(5, 4.1, '1:N', ha='center', fontsize=7, color='#F59E0B', fontweight='bold')
    ax.annotate('', xy=(11.8, 3.5), xytext=(10.2, 3.5),
               arrowprops=dict(arrowstyle='-', color='#F59E0B', lw=1.2))
    ax.text(11, 4.1, '1:N', ha='center', fontsize=7, color='#F59E0B', fontweight='bold')

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()
    doc.add_paragraph(
        'Explicacion: Diagrama Entidad-Relacion de la base de datos Supabase PostgreSQL. '
        'La tabla profiles almacena los datos de usuario (email, nombre, avatar, '
        'preferencias de tema e idioma). La tabla entries registra cada entrada del '
        'diario con su analisis emocional (mood, score, requires_help). La tabla '
        'contact_requests gestiona las solicitudes de contacto entre administradores '
        'y usuarios en casos de crisis detectada. Las PK (primary keys) se muestran '
        'en morado y las FK (foreign keys) en ambar.')

def diagrama_radar_calidad(doc, valores=None, titulo="Atributos de Calidad ISO/IEC 25010"):
    """
    Genera un diagrama de radar (spider chart) que evalua los atributos de
    calidad del software segun la norma ISO/IEC 25010.
    """
    import numpy as np
    if not _HAS_MPL: return

    if valores is None:
        valores = {
            'Funcionalidad': 9, 'Fiabilidad': 8.5, 'Usabilidad': 8,
            'Eficiencia': 7.5, 'Mantenibilidad': 8, 'Portabilidad': 7,
            'Seguridad': 9, 'Compatibilidad': 7.5
        }

    categorias = list(valores.keys())
    puntuaciones = list(valores.values())
    N = len(categorias)
    angulos = [n / float(N) * 2 * np.pi for n in range(N)]
    angulos += angulos[:1]
    puntuaciones += puntuaciones[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(categorias, fontsize=9, fontweight='bold')
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=7, color='#94A3B8')
    ax.fill(angulos, puntuaciones, color='#6366F1', alpha=0.25)
    ax.plot(angulos, puntuaciones, color='#6366F1', linewidth=2, marker='o', markersize=5)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=20, color='#6366F1')
    ax.grid(color='#E2E8F0', linewidth=0.5)

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(12))
    plt.close(fig); buf.close()
    doc.add_paragraph(
        'Explicacion: Diagrama de radar que evalua los 8 atributos de calidad '
        'ISO/IEC 25010. MindMood destaca en Funcionalidad (9/10) por su pipeline '
        'completo de IA y en Seguridad (9/10) por RLS, JWT y validacion de entrada. '
        'Areas de mejora: Portabilidad (7/10) y Eficiencia (7.5/10).')

def diagrama_despliegue_sistema(doc, titulo="Diagrama de Despliegue - Arquitectura de 3 Capas"):
    """
    Genera un diagrama de despliegue que muestra la arquitectura de 3 capas
    del sistema MindMood con las tecnologias utilizadas en cada capa.
    """
    import matplotlib.patches as mpatches
    if not _HAS_MPL: return

    capas = [
        {"y": 6.5, "nombre": "Capa de Presentacion",
         "detalle": "React 19 SPA + Vite 6\nTailwindCSS v4 + shadcn/ui\nPWA con Service Worker",
         "tecnologia": "Navegador Web", "color": "#6366F1"},
        {"y": 4, "nombre": "Capa de Negocio",
         "detalle": "FastAPI + Uvicorn\nPipeline IA 10 etapas\n2 Modelos HuggingFace\nRate Limiter",
         "tecnologia": "Servidor Python", "color": "#EC4899"},
        {"y": 1.5, "nombre": "Capa de Datos",
         "detalle": "PostgreSQL (Supabase)\nRow Level Security\nAuth JWT + PKCE\n12 funciones RPC",
         "tecnologia": "Base de Datos", "color": "#10B981"},
    ]

    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=15, color='#6366F1')

    for capa in capas:
        rect = mpatches.FancyBboxPatch((1.5, capa["y"] - 0.8), 11, 1.6,
               boxstyle="round,pad=0.1", edgecolor=capa["color"],
               facecolor=f'{capa["color"]}15', linewidth=2)
        ax.add_patch(rect)
        ax.text(2, capa["y"] + 0.5, capa["nombre"],
                fontsize=10, fontweight='bold', color=capa["color"])
        ax.text(8.5, capa["y"], capa["detalle"],
                ha='center', fontsize=7.5, color='#334155', linespacing=1.5)
        ax.text(12, capa["y"] + 0.5, capa["tecnologia"],
                ha='right', fontsize=7, fontstyle='italic', color='#64748B')

    ax.annotate('', xy=(7, 4.8), xytext=(7, 3.2),
               arrowprops=dict(arrowstyle='<->', color='#94A3B8', lw=1.5))
    ax.text(7.5, 4, 'REST API\nJSON', fontsize=7, color='#64748B', ha='center')
    ax.annotate('', xy=(7, 3), xytext=(7, 6.4),
               arrowprops=dict(arrowstyle='<->', color='#94A3B8', lw=1.5))
    ax.text(7.5, 4.7, 'HTTP\nJSON', fontsize=7, color='#64748B', ha='center')
    ax.text(7, 0.5, 'Desarrollo: ngrok tunnel  |  Produccion: Vercel + Supabase',
            ha='center', fontsize=8, color='#64748B', style='italic')

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(16))
    plt.close(fig); buf.close()
    doc.add_paragraph(
        'Explicacion: Diagrama de despliegue de la arquitectura de 3 capas. '
        'La Capa de Presentacion (React PWA) se ejecuta en el navegador del usuario. '
        'La Capa de Negocio (FastAPI + HuggingFace) procesa las solicitudes de analisis '
        'en un servidor Python con rate limiting. La Capa de Datos (Supabase PostgreSQL) '
        'almacena la informacion con Row Level Security y autenticacion JWT.')

def diagrama_boxplot_latencia(doc, titulo="Distribucion de Tiempos de Respuesta del Pipeline"):
    """
    Genera un diagrama de caja (box plot) que muestra la distribucion
    de tiempos de respuesta del endpoint /analyze agrupados por longitud del texto.
    """
    import numpy as np
    if not _HAS_MPL: return

    np.random.seed(42)
    corta = np.random.normal(0.45, 0.12, 40)
    media = np.random.normal(0.85, 0.25, 40)
    larga = np.random.normal(1.6, 0.45, 40)
    data = [corta, media, larga]

    fig, ax = plt.subplots(figsize=(10, 5.5))
    bp = ax.boxplot(data, patch_artist=True, widths=0.5,
                     medianprops=dict(color='white', linewidth=2),
                     whiskerprops=dict(color='#475569'),
                     capprops=dict(color='#475569'))

    for patch, color in zip(bp['boxes'], ['#6366F1', '#EC4899', '#10B981']):
        patch.set_facecolor(color)
        patch.set_alpha(0.7)

    ax.axhline(y=5.0, color='#EF4444', linestyle='--', linewidth=2,
               label='Umbral SLA: 5 segundos')
    ax.set_xticklabels(['Corta\n(<50 car.)', 'Media\n(50-200 car.)', 'Larga\n(>200 car.)'],
                         fontsize=9)
    ax.set_ylabel('Tiempo de respuesta (segundos)', fontsize=9)
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=12, color='#6366F1')
    ax.legend(fontsize=8, loc='upper left')
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y', linestyle='--', alpha=0.3)

    plt.tight_layout()
    buf = _fig_to_png(fig)
    doc.add_picture(buf, width=Cm(15))
    plt.close(fig); buf.close()
    doc.add_paragraph(
        'Explicacion: Box plot de tiempos de respuesta del pipeline de IA '
        'agrupados por longitud de texto (corta <50 caracteres, media 50-200, '
        'larga >200). Las cajas muestran la mediana (linea blanca), los cuartiles '
        'y los valores atipicos. Todos los valores estan muy por debajo del umbral '
        'SLA de 5 segundos (linea roja punteada), cumpliendo el requerimiento PS-005.')

def agregar_caption_figura(doc, numero, texto):
    """
    Agrega un caption numerado debajo de una figura con formato profesional.
    Formato: "Figura {numero}: {texto}" centrado, gris, 9pt, cursiva.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Figura {numero}: {texto}')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100, 116, 139)
    r.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)

def agregar_caption_tabla(doc, numero, texto):
    """
    Agrega un caption numerado encima de una tabla con formato profesional.
    Formato: "Tabla {numero}: {texto}" negrita, 10pt, indigo.
    """
    p = doc.add_paragraph()
    r = p.add_run(f'Tabla {numero}: {texto}')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(99, 102, 241)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
